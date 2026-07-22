from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from pathlib import Path


OUTPUT = Path("artifacts/quality-document-workflow-guide-th.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "59636E"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CAUTION = "FFF4D6"
CAUTION_TEXT = "7A5A00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    pPr.append(keep)


def set_font(run, size=11, bold=None, color=INK, italic=None):
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:ascii"), "TH Sarabun New")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "TH Sarabun New")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def clear_paragraph(p):
    p._element.clear_content()


def add_text(p, text, size=11, bold=None, color=INK, italic=None):
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    return run


def apply_style_tokens(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "TH Sarabun New"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "TH Sarabun New")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "TH Sarabun New")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        st = styles[name]
        st.font.name = "TH Sarabun New"
        st._element.rPr.rFonts.set(qn("w:ascii"), "TH Sarabun New")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "TH Sarabun New")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "TH Sarabun New"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def setup_section(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_font(run, size=9, color=MUTED)


def setup_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    add_text(p, "คู่มืออ้างอิง Workflow โมดูลเอกสารคุณภาพ", size=9, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, "เอกสารร่างเพื่อจัดทำ WI/QP | หน้า ", size=9, color=MUTED)
    add_page_number(p)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_text(p, text, size={1: 16, 2: 13, 3: 12}[level], bold=True, color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])
    set_keep_with_next(p)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True)
        add_text(p, text[len(bold_prefix):])
    else:
        add_text(p, text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_text(p, item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        add_text(p, item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, text, size=10.5, bold=True, color=DARK_BLUE)
    for i, values in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            if i % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            clear_paragraph(p)
            add_text(p, str(value), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, text, caution=False):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CAUTION if caution else LIGHT_GRAY)
    p = cell.paragraphs[0]
    clear_paragraph(p)
    color = CAUTION_TEXT if caution else DARK_BLUE
    add_text(p, title + " ", size=10.5, bold=True, color=color)
    add_text(p, text, size=10.5, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FC")
    p = cell.paragraphs[0]
    clear_paragraph(p)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_font(run, size=10, color=DARK_BLUE)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_document():
    doc = Document()
    section = doc.sections[0]
    setup_section(section)
    apply_style_tokens(doc)
    setup_header_footer(section)

    # First page - memo masthead adapted for a formal internal guide.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(3)
    add_text(p, "เอกสารร่างเพื่อจัดทำ WI/QP", size=12, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_text(p, "คู่มืออ้างอิง Workflow", size=24, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    add_text(p, "โมดูลเอกสารคุณภาพ", size=20, bold=True, color=DARK_BLUE)
    add_text(p, " (Quality Documents)", size=14, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, [2200, 7160])
    metadata = [
        ("วัตถุประสงค์", "ใช้เป็นเอกสารอ้างอิงสำหรับจัดทำ WI/QP การควบคุมเอกสารคุณภาพ"),
        ("ขอบเขต", "ตั้งแต่การสร้างเอกสาร ควบคุมไฟล์ Revision การเผยแพร่ การอ่าน การทบทวน และการยกเลิก"),
        ("ผู้ใช้งานหลัก", "Reviewer, Document Controller (DCC), Quality Manager, Laboratory Director และ Admin"),
        ("วันที่จัดทำ", "22 กรกฎาคม 2569 (ร่างจากพฤติกรรมที่ระบบรองรับปัจจุบัน)"),
    ]
    for idx, (label, value) in enumerate(metadata):
        c1, c2 = meta.rows[idx].cells
        set_cell_shading(c1, LIGHT_BLUE)
        for cell in (c1, c2):
            p = cell.paragraphs[0]
            clear_paragraph(p)
        add_text(c1.paragraphs[0], label, size=10.5, bold=True, color=DARK_BLUE)
        add_text(c2.paragraphs[0], value, size=10.5)

    doc.add_paragraph()
    add_callout(doc, "การใช้งานเอกสารนี้:", "เนื้อหานี้อธิบายสิ่งที่ระบบทำได้จริง ควรกำหนดผู้มีอำนาจอนุมัติ การเผยแพร่ การเก็บรักษา และกรณียกเว้นให้ชัดใน QP ก่อนประกาศใช้", caution=True)

    add_heading(doc, "สารบัญเนื้อหา", 1)
    toc_items = [
        "ภาพรวมและขอบเขตของโมดูล", "บทบาทและสิทธิ์", "คลังเอกสาร Dashboard Categories และ Master List",
        "การสร้างเอกสารใหม่และการควบคุมไฟล์", "Workflow สถานะเอกสาร", "Revision และประวัติเอกสาร",
        "การอัปโหลดเอกสารใหม่เป็นชุด: QP/WI พร้อม RF/FM/CF", "เอกสารที่เกี่ยวข้องและไฟล์แนบ",
        "การอ่าน ดาวน์โหลด และรายงานการอ่าน", "การทบทวนประจำปี", "การยกเลิก ลบ และการเก็บหลักฐาน",
        "ข้อกำหนดที่เสนอสำหรับ WI/QP",
    ]
    add_numbers(doc, toc_items)
    doc.add_page_break()

    add_heading(doc, "1. ภาพรวมและขอบเขตของโมดูล", 1)
    add_body(doc, "โมดูลเอกสารคุณภาพเป็นระบบควบคุมวงจรชีวิตเอกสาร ตั้งแต่การจัดทำฉบับร่าง การทบทวน การอนุมัติ การเผยแพร่ การแก้ไข Revision การติดตามการอ่าน การทบทวนประจำปี จนถึงการยกเลิกใช้งาน")
    add_table(doc, ["กลุ่มงาน", "หน้าที่หลัก"], [
        ("Document dashboard", "ติดตามจำนวนเอกสาร สถานะ งานค้าง และเอกสารใกล้ครบกำหนดทบทวน"),
        ("Library", "ค้นหา กรอง สร้าง แก้ไข เปิดอ่าน ดาวน์โหลด และจัดการเอกสาร"),
        ("Pending", "รวมคิวงานสำหรับ DCC/Reviewer เช่น รอทำ PDF รอ Review รอเผยแพร่ และรอ annual review"),
        ("Categories", "นำทางเอกสารตามหน่วยงานและประเภท"),
        ("Master List", "ทะเบียนเอกสารคุณภาพและการส่งออกสำหรับพิมพ์"),
        ("Read report", "ติดตามกลุ่มเป้าหมาย ผู้ที่อ่านแล้ว และผู้ที่ยังไม่อ่าน"),
    ], [2600, 6760])

    add_heading(doc, "2. ประเภทเอกสารและคำจำกัดความ", 1)
    add_table(doc, ["รหัส", "ประเภท", "ลักษณะการควบคุม"], [
        ("QM", "Quality Manual", "ติดตามรอบทบทวน; ใช้ Rev+ เมื่อมีการเปลี่ยนแปลง"),
        ("QP", "Quality Procedure", "เอกสารควบคุมหลัก; ต้องมี source Word/Excel และ PDF ก่อน Review"),
        ("WI", "Work Instruction", "เอกสารควบคุมหลัก; ต้องมี source Word/Excel และ PDF ก่อน Review"),
        ("FM / FR", "Form", "เอกสารแบบฟอร์ม; ใช้ status, revision และ history"),
        ("RF", "Reference", "เอกสารอ้างอิง; ใช้ status, revision และ history"),
        ("CF", "Card file", "เอกสาร Card file; ใช้ status, revision และ history"),
        ("Lb / Manual / Policy / Others", "เอกสารประเภทอื่น", "ใช้ workflow และ revision ตามความเหมาะสม; ไม่สร้าง cover เฉพาะประเภท"),
    ], [1500, 2500, 5360])
    add_body(doc, "คำว่า “ไฟล์ต้นฉบับ” หมายถึง Word/Excel ส่วน “ไฟล์ทางการ” หมายถึงไฟล์ที่ DCC รับรองให้ใช้เป็นเอกสารควบคุมในระบบ")

    add_heading(doc, "3. บทบาทและหน้าที่", 1)
    add_table(doc, ["บทบาท", "หน้าที่ที่ควรกำหนดใน WI/QP", "สิ่งที่ระบบรองรับ"], [
        ("Reviewer", "ทบทวนเนื้อหา ตรวจความถูกต้อง และยืนยัน annual review", "จัดทำ/แก้ไข draft, จัดการไฟล์แนบ, ยืนยัน annual review; ไม่มีสิทธิ์อัปโหลด PDF ทางการหรือเผยแพร่ตามสิทธิ์ workflow ปกติ"),
        ("DCC", "ควบคุมรหัส Metadata รูปแบบ ไฟล์ Revision และการเผยแพร่", "อัปโหลด/ยืนยันไฟล์ทางการ, ทำ PDF, จัดการคิว, annual review, set workflow, read audience"),
        ("Quality Manager / Laboratory Director", "พิจารณาและอนุมัติความถูกต้องด้านคุณภาพ/วิชาการ", "มีสิทธิ์ workflow สำหรับ Review → Approved และ Approved → Published ตามบทบาทที่กำหนด"),
        ("Admin", "กำกับสิทธิ์และจัดการกรณีพิเศษ", "จัดการได้ครบ รวมถึง import, purge และการดำเนินการพิเศษ"),
        ("Viewer", "อ่านเอกสารควบคุมที่ประกาศใช้", "เข้าถึง Published เท่านั้น"),
    ], [1500, 3900, 3960])
    add_callout(doc, "ข้อควรควบคุม:", "สิทธิ์เชิงเทคนิคของ DCC ในระบบกว้างกว่าหน้าที่ที่ควรแยกในระบบคุณภาพ จึงควรกำหนดว่าการอนุมัติเนื้อหาต้องมีหลักฐานจาก Quality Manager/Laboratory Director ก่อน DCC เผยแพร่", caution=True)

    add_heading(doc, "4. การสร้างเอกสารใหม่", 1)
    add_numbers(doc, [
        "ผู้จัดทำเลือกสร้างเอกสารใหม่จาก Library หรือ Dashboard ระบบสร้างเอกสารใหม่เป็นสถานะ Draft เท่านั้น",
        "กรอกข้อมูลควบคุม: รหัส ชื่อ ประเภท หน่วยงาน Revision ผู้จัดทำ ผู้รับรอง ผู้อนุมัติ วันที่แก้ไข/ทบทวน วันที่มีผลใช้ คำอธิบาย และการเผยแพร่ภายใน/สาธารณะ",
        "ระบบตรวจรหัสเอกสารซ้ำก่อนบันทึก เพื่อป้องกันการมีเอกสารรหัสเดียวกันในคลัง",
        "อัปโหลดไฟล์ต้นฉบับและ/หรือไฟล์ทางการตามประเภทเอกสาร",
        "สำหรับ QP/WI ให้เลือกอัปโหลดเอกสารเป็นชุดได้หลังบันทึก Draft หลักสำเร็จ",
    ])
    add_heading(doc, "4.1 การดึงข้อมูลเพื่อช่วยกรอก", 2)
    add_bullets(doc, [
        "หน้าสร้างเอกสารหลักรองรับการอ่าน PDF, DOCX และ XLSX เพื่อช่วยดึงข้อความและข้อมูล header",
        "จำกัดไฟล์สำหรับการดึงข้อมูลอัตโนมัติที่ 20 MB",
        "ผลการดึงข้อมูลเป็นข้อมูลช่วยกรอก ผู้จัดทำหรือ DCC ต้องตรวจทานรหัส ชื่อ Revision วันที่ และผู้รับผิดชอบก่อนบันทึก",
        "การอัปโหลดเอกสารเป็นชุดจะดึงรหัสและชื่อจากชื่อไฟล์เป็นหลัก ไม่ได้อ่านเนื้อหา RF/FM/CF แต่ละไฟล์อัตโนมัติ",
    ])

    add_heading(doc, "5. การควบคุมไฟล์", 1)
    add_table(doc, ["ประเภทไฟล์", "QP/WI", "ประเภทอื่น"], [
        ("ไฟล์ต้นฉบับ", "DOC, DOCX, XLS, XLSX; จำเป็นก่อนเข้า Review", "อัปโหลดได้ตามความจำเป็น"),
        ("ไฟล์ทางการ", "PDF เนื้อหา; DCC/Admin เป็นผู้มีสิทธิ์อัปโหลด", "PDF, DOC, DOCX, XLS, XLSX ตามความเหมาะสม"),
        ("ขนาดสูงสุด", "50 MB ต่อไฟล์", "50 MB ต่อไฟล์"),
        ("การยืนยันไฟล์", "DCC ตรวจไฟล์ทางการก่อนเลื่อน workflow", "ไฟล์จากผู้จัดทำอาจอยู่ pending file จน DCC ยืนยัน"),
    ], [2400, 3480, 3480])
    add_callout(doc, "การตั้งค่าปัจจุบัน:", "ระบบระงับการสร้างหน้าปก QP/WI อัตโนมัติ ดังนั้น PDF ที่ DCC อัปโหลดและเผยแพร่เป็นไฟล์ควบคุมจริง หากหน่วยงานต้องใช้หน้าปกจากระบบ ต้องเปิดใช้ฟังก์ชันและทดสอบก่อนประกาศใช้", caution=True)

    add_heading(doc, "6. Workflow สถานะเอกสาร", 1)
    add_code_block(doc, ["Draft  →  Review  →  Approved  →  Published  →  Obsolete", "             ↖             ↖", "          กลับ Draft      กลับ Review"])
    add_table(doc, ["สถานะ", "วัตถุประสงค์", "ผลการควบคุม"], [
        ("Draft", "จัดทำและแก้ไขเนื้อหา", "แก้ไข metadata และไฟล์ได้ตามสิทธิ์; QP/WI ต้องมี source และ PDF ก่อน Review"),
        ("Review", "พร้อมให้ทบทวน/อนุมัติ", "ส่งกลับ Draft เพื่อแก้ไขได้"),
        ("Approved", "เนื้อหาอนุมัติแล้ว", "ตรวจความพร้อมก่อนประกาศใช้; ส่งกลับ Review ได้"),
        ("Published", "ฉบับควบคุมที่ใช้งานจริง", "เนื้อหาและ Revision ต้องแก้ผ่าน working revision เท่านั้น"),
        ("Obsolete", "ยกเลิกใช้งาน", "PDF ถูกประทับ OBSOLETE เมื่อทำได้; ถือเป็นสถานะสิ้นสุด"),
    ], [1600, 2800, 4960])
    add_heading(doc, "6.1 ขั้นตอนมาตรฐาน", 2)
    add_numbers(doc, [
        "Reviewer/ผู้จัดทำจัดทำเนื้อหาและส่งไฟล์ต้นฉบับใน Draft",
        "DCC ตรวจข้อมูลควบคุม จัดทำ/ตรวจ PDF และยืนยันไฟล์ทางการ",
        "DCC ส่ง Draft เข้า Review เมื่อไฟล์และข้อมูลครบ",
        "ผู้มีอำนาจอนุมัติพิจารณา Review แล้วเปลี่ยนเป็น Approved หรือส่งกลับแก้ไข",
        "DCC ตรวจขั้นสุดท้ายและเผยแพร่ตามอำนาจและหลักฐานการอนุมัติที่กำหนด",
        "ระบบบันทึกสถานะ วันเวลา และผู้ดำเนินการใน status history/audit log",
    ])

    add_heading(doc, "7. Revision และประวัติเอกสาร", 1)
    add_code_block(doc, ["Published Rev. เดิม", "        ↓", "Working revision (Rev. ใหม่) → Draft → Review → Approved → Published", "        ↓", "ฉบับเดิมถูกเก็บใน Revision History"])
    add_bullets(doc, [
        "สร้าง working revision ได้เฉพาะเอกสาร Published และมีได้เพียง 1 working revision ที่ยังไม่จบต่อเอกสาร",
        "ฉบับ Published เดิมยังคงเป็นฉบับที่ใช้งานได้ ระหว่างที่ Rev. ใหม่อยู่ระหว่างดำเนินการ",
        "เมื่อ Rev. ใหม่ Published ระบบเก็บฉบับเดิมเป็น Revision History และเลื่อนฉบับใหม่เป็น current document",
        "สามารถยกเลิก working revision ที่ยังไม่เผยแพร่ได้",
        "Quick Update (Upd+) ใช้สำหรับเอกสาร non-controlled เช่น Form, Reference, Card file, Lb, Policy และ Others เพื่อเปลี่ยนไฟล์และเพิ่ม Revision แบบย่อ",
        "DCC/Admin เพิ่มประวัติ Revision ย้อนหลังจากระบบเดิมได้ โดยไม่กระทบฉบับปัจจุบัน",
    ])
    add_callout(doc, "ความเสี่ยง:", "ระบบมีฟังก์ชันย้อน current revision ให้ DCC/Admin ใช้ฉบับก่อนหน้าได้ ควรกำหนดเป็นกรณีฉุกเฉิน ต้องมีการอนุมัติและบันทึกเหตุผล ไม่ควรใช้แทน Rev+ ปกติ", caution=True)

    add_heading(doc, "8. การอัปโหลดเอกสารใหม่เป็นชุด: QP/WI พร้อม RF/FM/CF", 1)
    add_body(doc, "ใช้เมื่อ QP หรือ WI ฉบับใหม่ออกพร้อมเอกสารสนับสนุน เช่น Reference, Form และ Card file ซึ่งต้องควบคุมสถานะและ Revision ให้สัมพันธ์กับเอกสารหลัก")
    add_code_block(doc, ["สร้าง QP/WI หลักเป็น Draft", "        ↓", "เลือกอัปโหลดเอกสารเป็นชุด", "        ↓", "ลงทะเบียน RF/FM/CF ใหม่ | ลิงก์ฉบับ Published | เปิด Rev+ | แนบไฟล์", "        ↓", "DCC ตรวจไฟล์และความพร้อม", "        ↓", "Review ทั้งชุด → Approved ทั้งชุด → Published ทั้งชุด"])
    add_heading(doc, "8.1 เงื่อนไขเริ่มต้น", 2)
    add_bullets(doc, [
        "ใช้ได้กับ QP/WI ใหม่ที่เพิ่งสร้างใน Draft เท่านั้น ไม่ใช้กับการแก้ไขเอกสารเดิมหรือการ Import current document",
        "เพิ่มได้สูงสุด 30 ไฟล์ต่อครั้ง และขนาดไฟล์ละไม่เกิน 50 MB",
        "ผู้ใช้ต้องตรวจรายการแต่ละไฟล์ก่อนกดยืนยัน เพราะระบบจะดำเนินการตามประเภทที่เลือกทีละรายการ",
    ])
    add_heading(doc, "8.2 มาตรฐานการตั้งชื่อ RF/FM/CF", 2)
    add_code_block(doc, [
        "[ประเภท]-[รหัสเอกสารแม่]-[หน่วยงาน]-[เลขเอกสารย่อย] [ชื่อเอกสาร].[นามสกุล]",
        "FM-QP-LAB-03-05 แบบบันทึกการตรวจสอบ.pdf",
        "RF-QP-LAB-03 เอกสารอ้างอิง.pdf",
        "CF-QP-LAB-03-01 Card file.pdf",
    ])
    add_table(doc, ["คำนำหน้าชื่อไฟล์", "ระบบตั้งประเภท", "ข้อมูลที่ดึงจากชื่อไฟล์"], [
        ("FM- หรือ FR-", "Form", "คำก่อนช่องว่างเป็นรหัส; ข้อความหลังรหัสเป็นชื่อเอกสาร"),
        ("RF-", "Reference", "คำก่อนช่องว่างเป็นรหัส; ข้อความหลังรหัสเป็นชื่อเอกสาร"),
        ("CF-", "Card file", "คำก่อนช่องว่างเป็นรหัส; ข้อความหลังรหัสเป็นชื่อเอกสาร"),
        ("ชื่ออื่น", "ไฟล์แนบโดยค่าเริ่มต้น", "ผู้ใช้สามารถเปลี่ยนให้ลงทะเบียนเป็นเอกสารได้"),
    ], [2400, 2500, 4460])
    add_body(doc, "ระบบไม่บังคับตัวพิมพ์ใหญ่/เล็กในชื่อไฟล์ แต่เมื่อบันทึกจะใช้รหัสตัวพิมพ์ใหญ่ จึงควรใช้ FM-, RF- และ CF- ให้สม่ำเสมอ")
    add_heading(doc, "8.3 การจัดการเมื่อพบรหัสซ้ำ", 2)
    add_table(doc, ["สถานะที่พบ", "วิธีดำเนินการ"], [
        ("ไม่พบรหัส", "ลงทะเบียนเป็นเอกสารใหม่และผูกเป็นสมาชิกของชุด"),
        ("พบ Published", "เลือก “ลิงก์เอกสารเดิม” หรือ “เปิด Rev+ เอกสารเดิม”"),
        ("พบ Draft/Review/Approved", "ไม่สามารถลงทะเบียนซ้ำ ต้องแก้ไขหรือรอให้ workflow เดิมเสร็จ"),
        ("พบรหัสซ้ำภายในไฟล์ที่เลือก", "ระบบไม่ให้ยืนยันจนกว่าจะแก้รหัสให้ไม่ซ้ำ"),
    ], [3000, 6360])
    add_heading(doc, "8.4 ผลของทางเลือกแต่ละรายการ", 2)
    add_table(doc, ["ทางเลือก", "ผลที่ระบบบันทึก", "ผลต่อ workflow ชุด"], [
        ("ลงทะเบียนเอกสารใหม่", "สร้าง RF/FM/CF เป็น Draft ใหม่ พร้อมลิงก์แบบ registered", "ต้องเดินสถานะพร้อมเอกสารหลัก"),
        ("ลิงก์เอกสารเดิม", "เชื่อมเอกสาร Published แบบ linked", "ไม่เปลี่ยนสถานะเอกสารเดิม"),
        ("เปิด Rev+", "สร้าง working revision ของเอกสาร Published แบบ revision", "working revision ต้องเดินสถานะพร้อมเอกสารหลัก"),
        ("แนบไฟล์", "เก็บเป็น attachment ชั่วคราว/ประกอบของเอกสารหลัก", "ไม่ใช่เอกสารควบคุมแยกฉบับ"),
    ], [2400, 4000, 2960])
    add_heading(doc, "8.5 ขั้นตอนดำเนินงานของชุดเอกสาร", 2)
    add_numbers(doc, [
        "ผู้จัดทำสร้าง QP/WI หลักเป็น Draft และเปิดหน้าการอัปโหลดเอกสารเป็นชุด",
        "เลือกไฟล์ ระบบแยก FM/FR/RF/CF เป็นรายการลงทะเบียนอัตโนมัติ และไฟล์อื่นเป็น attachment",
        "ตรวจรหัส ชื่อ ประเภท หน่วยงาน และตัวเลือกเมื่อพบรหัสซ้ำ",
        "ยืนยันรายการ ระบบอัปโหลดและบันทึกผลสำเร็จ/ไม่สำเร็จเป็นรายไฟล์",
        "DCC ตรวจไฟล์ทางการและความครบถ้วนของทุกเอกสารที่ต้องเดิน workflow",
        "ใช้คำสั่งส่งทั้งชุดเข้า Review, อนุมัติทั้งชุด และเผยแพร่ทั้งชุดตามลำดับ",
        "DCC ตรวจผลหลังดำเนินการทุกครั้ง และแก้ไขรายการที่ล้มเหลวก่อนดำเนินการขั้นถัดไป",
    ])
    add_callout(doc, "ข้อควรระวัง:", "การเลื่อนสถานะชุดทำทีละรายการตามลำดับ จึงอาจเกิดกรณีบางรายการสำเร็จและบางรายการล้มเหลวได้ DCC ต้องตรวจผลและทำให้สถานะเอกสารหลักกับสมาชิกที่ต้องเดิน workflow กลับมาสอดคล้องกันก่อนเผยแพร่", caution=True)

    add_heading(doc, "9. เอกสารที่เกี่ยวข้องและไฟล์แนบ", 1)
    add_table(doc, ["ฟังก์ชัน", "ใช้เมื่อ", "ผลต่อเอกสาร"], [
        ("Related documents", "ต้องการอ้างอิง QP/WI/Form/Reference ที่เกี่ยวข้อง", "สร้างความสัมพันธ์เพื่อการอ้างอิง แต่ไม่บังคับให้ workflow เดินพร้อมกัน"),
        ("Attachments", "มีหลักฐาน supporting document หรือไฟล์ประกอบ", "เก็บไฟล์แนบหลายไฟล์; ไม่สร้างเลขเอกสารควบคุมใหม่"),
        ("Download ZIP", "ต้องส่งมอบไฟล์หลัก ต้นฉบับ ลิงก์ และไฟล์แนบพร้อมกัน", "รวบรวมตามสิทธิ์ของผู้ใช้และบันทึกการเข้าถึงเมื่อเกี่ยวข้อง"),
    ], [2400, 3200, 3760])

    add_heading(doc, "10. การอ่าน ดาวน์โหลด และรายงานการอ่าน", 1)
    add_bullets(doc, [
        "เมื่อผู้ใช้เปิดอ่าน ระบบบันทึก view log และสร้างลิงก์ชั่วคราวสำหรับเข้าถึงไฟล์",
        "Viewer เปิดได้เฉพาะเอกสาร Published",
        "ผู้มีสิทธิ์ดาวน์โหลดไฟล์หลัก ไฟล์ต้นฉบับ ไฟล์แนบ Revision และชุดเอกสารเป็น ZIP ได้ตามบทบาท",
        "Read report ใช้ติดตาม QM/QP/WI/Manual ที่ Published โดยนับผู้อ่านเฉพาะหลังวันที่เผยแพร่ของฉบับปัจจุบัน",
        "DCC/Admin กำหนดกลุ่มเป้าหมายผู้ต้องอ่านได้ทั้งตามหน่วยงานและรายบุคคล",
        "เมื่อเผยแพร่ Rev. ใหม่ ยอดการอ่านเริ่มนับตาม published date ใหม่; annual review แบบไม่มีการแก้ไขไม่ reset ยอดอ่าน",
    ])

    add_heading(doc, "11. การทบทวนประจำปี", 1)
    add_body(doc, "ระบบติดตามรอบทบทวนของ QM, QP, WI และ Manual โดยคำนวณครบกำหนด 1 ปีจากวันที่ล่าสุดระหว่างวันที่แก้ไข วันที่ทบทวนล่าสุด และวันที่ทบทวนเดิม ระบบเริ่มแสดงการเตือน 90 วันก่อนครบกำหนด")
    add_table(doc, ["ประเภท", "แนวทางเมื่อไม่มีการแก้ไข"], [
        ("QP / WI", "Reviewer/DCC/Admin ยืนยัน “ทบทวนแล้ว” จากนั้น DCC ดำเนินการแบบกลุ่ม ระบบเพิ่มประวัติ Rev. '-' และอัปเดต last reviewed date โดยไม่เปลี่ยน Revision/เนื้อหา/วันที่มีผลใช้"),
        ("QM / Manual", "ใช้ Rev+ ปกติ แม้ไม่มีการแก้ไข เพราะไม่มี review-only workflow"),
    ], [1900, 7460])

    add_heading(doc, "12. การยกเลิก ลบ และการเก็บหลักฐาน", 1)
    add_heading(doc, "12.1 Obsolete", 2)
    add_numbers(doc, [
        "DCC ตรวจว่ามีเอกสารทดแทนหรือมาตรการป้องกันการใช้ฉบับเดิมแล้ว",
        "บันทึกเหตุผลและวันที่ยกเลิก",
        "เปลี่ยน Published เป็น Obsolete",
        "ระบบประทับ watermark OBSOLETE บน PDF เมื่อดำเนินการได้ และบันทึก audit log",
    ])
    add_heading(doc, "12.2 การลบ", 2)
    add_bullets(doc, [
        "การลบปกติเป็น soft delete; เอกสารไม่แสดงในคลังปกติแต่ยังอยู่ในระบบ",
        "Admin/Manager สามารถ purge เอกสารที่ soft delete พร้อมลบไฟล์จาก storage แบบถาวร",
        "ก่อน purge ต้องกำหนด retention period และผู้อนุมัติใน QP เพื่อไม่ให้ลบหลักฐานคุณภาพก่อนกำหนด",
    ])
    add_heading(doc, "12.3 Audit trail", 2)
    add_table(doc, ["หลักฐาน", "ตัวอย่างข้อมูลที่เก็บ"], [
        ("Access log", "upload, download, edit, delete, view"),
        ("Status history", "สถานะปลายทาง วันที่เวลา และผู้ดำเนินการ"),
        ("Revision history", "Revision เดิม เหตุผลแก้ไข ผู้แก้ไข/ผู้อนุมัติ และไฟล์เดิม"),
        ("Audit log", "import, create draft, official confirm, annual review, obsolete stamp"),
        ("Read log", "ผู้อ่านและเวลาอ่าน"),
    ], [2700, 6660])

    add_heading(doc, "13. ข้อกำหนดที่เสนอสำหรับนำไปเขียน WI/QP", 1)
    add_numbers(doc, [
        "ระบุเจ้าของกระบวนการ: ผู้จัดทำ, Reviewer, DCC, ผู้อนุมัติ และ Admin",
        "กำหนดเกณฑ์อนุมัติเนื้อหาให้แยกจากการควบคุมรูปแบบและการเผยแพร่ของ DCC",
        "กำหนดมาตรฐานรหัสและชื่อไฟล์ โดยเฉพาะ FM/FR/RF/CF ที่ใช้กับการอัปโหลดเป็นชุด",
        "กำหนดว่ากรณีใดต้องใช้ Document Set และกรณีใดเป็นเพียง attachment",
        "กำหนดหลักฐานที่ต้องมีในแต่ละสถานะ รวมถึง source, PDF ทางการ, เหตุผลแก้ไข และการอนุมัติ",
        "กำหนดการจัดการ partial failure ของชุดเอกสาร และการตรวจสอบความสอดคล้องของสถานะหลังทำรายการ",
        "กำหนดความถี่ทบทวนประจำปี วิธีจัดการ QM/Manual และหลักฐานการทบทวน",
        "กำหนด retention period, การ obsolete, soft delete, purge และการใช้งาน rollback ในกรณีพิเศษ",
        "กำหนดเกณฑ์การติดตามการอ่าน และการดำเนินการเมื่อบุคลากรไม่อ่านเอกสารภายในเวลาที่กำหนด",
    ])

    add_heading(doc, "ภาคผนวก A: Checklist ของ DCC ก่อนเผยแพร่", 1)
    checklist = [
        "ตรวจรหัสเอกสาร ประเภท ชื่อ หน่วยงาน และ Revision ถูกต้อง",
        "ตรวจชื่อผู้จัดทำ ผู้รับรอง ผู้อนุมัติ และวันที่มีผลใช้ครบถ้วน",
        "ตรวจว่า QP/WI มี source Word/Excel และ PDF ทางการครบ",
        "ตรวจว่าเอกสารในชุด RF/FM/CF มีสถานะและ Revision สอดคล้องกับเอกสารหลัก",
        "ตรวจลิงก์เอกสาร Published และไฟล์แนบว่าถูกต้อง",
        "ตรวจกลุ่มผู้ต้องอ่านเอกสาร",
        "ตรวจผลการอนุมัติและหลักฐานก่อนกด Published",
        "ตรวจว่า Master List, Revision History, status history และ audit trail ถูกบันทึกหลังเผยแพร่",
    ]
    add_table(doc, ["ลำดับ", "รายการตรวจสอบ", "ผลตรวจ/หมายเหตุ"], [(i + 1, item, "") for i, item in enumerate(checklist)], [900, 5700, 2760])

    add_heading(doc, "ภาคผนวก B: Checklist การอัปโหลด QP/WI พร้อม RF/FM/CF", 1)
    set_checklist = [
        "QP/WI หลักถูกสร้างเป็น Draft แล้ว",
        "ไฟล์ FM/FR/RF/CF ตั้งชื่อในรูปแบบที่กำหนด และมีช่องว่างก่อนชื่อเอกสาร",
        "ตรวจชนิดเอกสารที่ระบบเสนอ: FM/FR=Form, RF=Reference, CF=Card file",
        "ตรวจรหัสซ้ำและเลือก Link existing หรือ Rev+ เมื่อจำเป็น",
        "แยกไฟล์ที่เป็นหลักฐานออกเป็น attachment ไม่ใช่เอกสารควบคุม",
        "ตรวจผลการอัปโหลดรายไฟล์และทำซ้ำเฉพาะรายการที่ไม่สำเร็จ",
        "DCC ตรวจไฟล์ทางการและความพร้อมของทุกสมาชิกก่อนส่งทั้งชุดเข้า Review",
        "ตรวจความสอดคล้องของสถานะทุกสมาชิกหลัง Review, Approved และ Published",
    ]
    add_table(doc, ["ลำดับ", "รายการตรวจสอบ", "ผลตรวจ/หมายเหตุ"], [(i + 1, item, "") for i, item in enumerate(set_checklist)], [900, 5700, 2760])

    doc.core_properties.title = "คู่มืออ้างอิง Workflow โมดูลเอกสารคุณภาพ"
    doc.core_properties.subject = "Workflow reference for Quality Documents module"
    doc.core_properties.author = "Lab Management Portal"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT.resolve())
